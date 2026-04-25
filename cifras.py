# -*- coding: utf-8 -*-
import streamlit as st
import requests
from bs4 import BeautifulSoup
from fpdf import FPDF
import re
import io
import zipfile
from docx import Document
from docx.shared import Pt

# -------------------------------------------------
# CONFIG
# -------------------------------------------------
st.set_page_config(
    page_title="Music Book Pro",
    page_icon="🎸",
    layout="wide"
)

st.markdown("""
<style>
textarea, pre, .cifra-renderizada{
    font-family:'Courier New', monospace !important;
    white-space:pre !important;
    word-wrap:normal !important;
    line-height:1 !important;
}

.stButton > button{
    width:100%;
    border-radius:5px;
}

.page-container{
    background:#0e1117;
    padding:20px;
    border-radius:10px;
    border:1px solid #333;
    color:white;
    margin-bottom:20px;
}

.cifra-renderizada{
    font-family:'Courier New', monospace !important;
    font-size:inherit !important;
    line-height:1 !important;
    margin:0 !important;
    padding:0 !important;
    display:block;
}

div[data-testid="stMarkdownContainer"] p{
    margin:0 !important;
    padding:0 !important;
}
</style>
""", unsafe_allow_html=True)

# -------------------------------------------------
# SESSION STATE
# -------------------------------------------------
if "book" not in st.session_state:
    st.session_state.book = []

if "limpador" not in st.session_state:
    st.session_state.limpador = 0

if "musica_focada" not in st.session_state:
    st.session_state.musica_focada = None

if "temp_titulo" not in st.session_state:
    st.session_state.temp_titulo = ""

if "temp_conteudo" not in st.session_state:
    st.session_state.temp_conteudo = ""

if "pdf_bytes" not in st.session_state:
    st.session_state.pdf_bytes = None

if "doc_bytes" not in st.session_state:
    st.session_state.doc_bytes = None

if "txt_bytes" not in st.session_state:
    st.session_state.txt_bytes = None

if "kindle_bytes" not in st.session_state:
    st.session_state.kindle_bytes = None

if "zip_bytes" not in st.session_state:
    st.session_state.zip_bytes = None

# -------------------------------------------------
# TRANSPOSIÇÃO
# -------------------------------------------------
NOTAS = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']

def transpor_acorde(acorde, semitons):
    def sub(match):
        nota = match.group(1)
        resto = match.group(2)
        
        if nota in NOTAS:
            idx = (NOTAS.index(nota) + semitons) % 12
            return NOTAS[idx] + resto
        
        return match.group(0)
    
    return re.sub(r'([A-G]#?)([^A-G\s]*)', sub, acorde)

def processar_texto(texto, semitons, colunas):
    if not texto:
        return ""
    
    linhas = texto.split("\n")
    linhas_t = []
    
    for linha in linhas:
        nova = ""
        pos = 0
        
        for m in re.finditer(r'\S+', linha):
            nova += " " * (m.start() - pos) + transpor_acorde(m.group(), semitons)
            pos = m.start() + len(m.group())
        
        linhas_t.append(nova + " " * (len(linha) - pos))
    
    if colunas == "2 Colunas":
        total = len(linhas_t)
        meio = (total // 2) + (total % 2)
        
        esq = linhas_t[:meio]
        dir = linhas_t[meio:]
        
        largura = max(len(x) for x in esq) if esq else 0
        
        final = []
        for i in range(max(len(esq), len(dir))):
            a = esq[i] if i < len(esq) else ""
            b = dir[i] if i < len(dir) else ""
            final.append(a.ljust(largura + 8) + b)
        
        return "\n".join(final)
    
    return "\n".join(linhas_t)

# -------------------------------------------------
# FUNÇÕES DE AJUSTE
# -------------------------------------------------
def ajustar_fonte(delta):
    if st.session_state.musica_focada is not None:
        idx = st.session_state.musica_focada
        atual = st.session_state.book[idx]["fonte"]
        novo = max(8, min(32, atual + delta))
        st.session_state.book[idx]["fonte"] = novo

def set_fonte(v):
    if st.session_state.musica_focada is not None:
        st.session_state.book[st.session_state.musica_focada]["fonte"] = v

def ajustar_tom(delta):
    if st.session_state.musica_focada is not None:
        st.session_state.book[st.session_state.musica_focada]["tom"] += delta

def set_tom(v):
    if st.session_state.musica_focada is not None:
        st.session_state.book[st.session_state.musica_focada]["tom"] = v

def set_layout(v):
    if st.session_state.musica_focada is not None:
        st.session_state.book[st.session_state.musica_focada]["cols"] = v

# -------------------------------------------------
# FUNÇÕES DE EXPORTAÇÃO
# -------------------------------------------------
def exportar_txt():
    linhas = []
    for m in st.session_state.book:
        linhas.append(f"\n{'='*50}\n{m['titulo'].upper()}\n{'='*50}\n")
        texto = processar_texto(m["conteudo"], m["tom"], m["cols"])
        linhas.append(texto)
    
    conteudo = "\n".join(linhas)
    return conteudo.encode("utf-8")

def exportar_doc():
    doc = Document()
    
    for m in st.session_state.book:
        doc.add_heading(m["titulo"], level=1)
        texto = processar_texto(m["conteudo"], m["tom"], m["cols"])
        paragrafo = doc.add_paragraph()
        run = paragrafo.add_run(texto)
        run.font.name = "Courier New"
        run.font.size = Pt(m["fonte"])
        doc.add_page_break()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def exportar_txt_simples():
    """Exporta em formato compatível com Kindle"""
    linhas = []
    for m in st.session_state.book:
        linhas.append(f"\n\n{'='*50}\n{m['titulo'].upper()}\n{'='*50}\n\n")
        texto = processar_texto(m["conteudo"], m["tom"], "1 Coluna")
        linhas.append(texto)
    
    return "\n".join(linhas).encode("utf-8")

def exportar_pdf_buffer():
    """Retorna o PDF em bytes"""
    pdf = FPDF()
    
    for m in st.session_state.book:
        pdf.add_page()
        pdf.set_font("Courier", "B", 14)
        pdf.cell(0, 10, m["titulo"], ln=True)
        pdf.set_font("Courier", size=m["fonte"])
        txt = processar_texto(m["conteudo"], m["tom"], m["cols"])
        
        # Converte para latin-1 (evita erro de encoding)
        try:
            txt = txt.encode('latin-1', 'replace').decode('latin-1')
        except:
            pass
        
        pdf.multi_cell(0, 5, txt)
    
    return bytes(pdf.output(dest="S"))

def exportar_completo_zip():
    """Exporta todos os formatos em um único ZIP"""
    buffer_zip = io.BytesIO()
    
    with zipfile.ZipFile(buffer_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
        pdf_bytes = exportar_pdf_buffer()
        zipf.writestr("repertorio.pdf", pdf_bytes)
        
        doc_bytes = exportar_doc()
        zipf.writestr("repertorio.docx", doc_bytes)
        
        txt_bytes = exportar_txt()
        zipf.writestr("repertorio.txt", txt_bytes)
        
        kindle_bytes = exportar_txt_simples()
        zipf.writestr("repertorio_kindle.txt", kindle_bytes)
    
    return buffer_zip.getvalue()

# -------------------------------------------------
# SIDEBAR
# -------------------------------------------------
st.sidebar.title("🎵 Music Book")

aba = st.sidebar.radio(
    "Navegação",
    ["Adicionar Música", "Visualizar Book", "Exportar"]
)

st.sidebar.divider()
st.sidebar.markdown("### 🛠️ Ajustes da Música Selecionada")

if st.session_state.musica_focada is not None:
    st.sidebar.write("Tamanho da Letra")
    c1, c2, c3 = st.sidebar.columns(3)
    
    if c1.button("A-", key="font_down"):
        ajustar_fonte(-1)
        st.rerun()
    
    if c2.button("11", key="font_reset"):
        set_fonte(11)
        st.rerun()
    
    if c3.button("A+", key="font_up"):
        ajustar_fonte(1)
        st.rerun()
    
    st.sidebar.write("Tom")
    t1, t2, t3 = st.sidebar.columns(3)
    
    if t1.button("♭", key="tom_down"):
        ajustar_tom(-1)
        st.rerun()
    
    if t2.button("0", key="tom_reset"):
        set_tom(0)
        st.rerun()
    
    if t3.button("♯", key="tom_up"):
        ajustar_tom(1)
        st.rerun()
    
    st.sidebar.write("Layout")
    l1, l2 = st.sidebar.columns(2)
    
    if l1.button("📄 1 Col", key="col1"):
        set_layout("1 Coluna")
        st.rerun()
    
    if l2.button("✂️ 2 Col", key="col2"):
        set_layout("2 Colunas")
        st.rerun()
else:
    st.sidebar.info("👈 Selecione uma música no Visualizar Book")

# -------------------------------------------------
# ADICIONAR MÚSICA
# -------------------------------------------------
if aba == "Adicionar Música":
    st.header("🔍 Capturar Cifra")
    
    url = st.text_input("Link da cifra (CifraClub)", key=f"url_{st.session_state.limpador}")
    
    if st.button("Capturar", key="capturar"):
        try:
            res = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
            soup = BeautifulSoup(res.text, "html.parser")
            
            titulo_elem = soup.find("h1", class_="t1") or soup.find("h1")
            titulo = titulo_elem.get_text().strip() if titulo_elem else "Sem título"
            
            pre_elem = soup.find("pre")
            cifra = pre_elem.get_text() if pre_elem else "Cifra não encontrada"
            
            st.session_state.temp_titulo = titulo
            st.session_state.temp_conteudo = cifra
            
            st.success("Cifra capturada com sucesso!")
            st.rerun()
        except Exception as e:
            st.error(f"Erro ao capturar cifra: {str(e)}")
    
    st.divider()
    
    tit = st.text_input("Título", value=st.session_state.temp_titulo, key="titulo_input")
    cif = st.text_area("Cifra", value=st.session_state.temp_conteudo, height=300, key="cifra_input")
    
    if st.button("✅ Adicionar ao Repertório", key="adicionar"):
        if tit and cif:
            st.session_state.book.append({
                "titulo": tit,
                "conteudo": cif,
                "fonte": 12,
                "tom": 0,
                "cols": "1 Coluna"
            })
            
            st.session_state.temp_titulo = ""
            st.session_state.temp_conteudo = ""
            st.session_state.limpador += 1
            
            st.success(f"Música '{tit}' adicionada com sucesso!")
            st.rerun()
        else:
            st.error("Preencha o título e a cifra!")

# -------------------------------------------------
# VISUALIZAR BOOK
# -------------------------------------------------
elif aba == "Visualizar Book":
    st.header("📖 Meu Repertório")
    
    if not st.session_state.book:
        st.info("O book está vazio. Adicione músicas na aba 'Adicionar Música'.")
    
    else:
        for i, m in enumerate(st.session_state.book):
            expander_label = f"🎸 {m['titulo']} (fonte: {m['fonte']}pt | tom: {m['tom']})"
            
            with st.expander(expander_label, expanded=(st.session_state.musica_focada == i)):
                
                # Botão para focar nesta música
                if st.button(f"🎯 Selecionar", key=f"select_{i}"):
                    st.session_state.musica_focada = i
                    st.rerun()
                
                texto_proc = processar_texto(m["conteudo"], m["tom"], m["cols"])
                
                st.markdown(
                    f"""
                    <div class='page-container'>
                    <div style='font-family:Courier New; font-size:{m["fonte"]}pt; line-height:1.2;'>
                    <h3>{m['titulo']}</h3>
                    """,
                    unsafe_allow_html=True
                )
                
                bloco = ""
                for linha in texto_proc.split("\n"):
                    if linha.strip() == "":
                        linha = " "
                    linha = linha.replace(" ", "&nbsp;")
                    bloco += f"<div class='cifra-renderizada'>{linha}</div>"
                
                st.markdown(bloco, unsafe_allow_html=True)
                st.markdown("</div></div>", unsafe_allow_html=True)
                
                if st.button(f"🗑️ Excluir Música", key=f"del_{i}"):
                    st.session_state.book.pop(i)
                    if st.session_state.musica_focada == i:
                        st.session_state.musica_focada = None
                    st.rerun()

# -------------------------------------------------
# EXPORTAR
# -------------------------------------------------
elif aba == "Exportar":
    st.header("📂 Exportar Livro")
    
    if not st.session_state.book:
        st.info("O book está vazio. Adicione músicas primeiro.")
    
    else:
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Formatos Individuais")
            
            # PDF
            if st.button("📑 Gerar PDF", key="gerar_pdf"):
                with st.spinner("Gerando PDF..."):
                    st.session_state.pdf_bytes = exportar_pdf_buffer()
            
            if st.session_state.pdf_bytes:
                st.download_button(
                    "📥 Baixar PDF",
                    data=st.session_state.pdf_bytes,
                    file_name="MeuRepertorio.pdf",
                    mime="application/pdf",
                    key="download_pdf"
                )
            
            # DOC
            if st.button("📝 Gerar DOC (Word)", key="gerar_doc"):
                with st.spinner("Gerando documento Word..."):
                    st.session_state.doc_bytes = exportar_doc()
            
            if st.session_state.doc_bytes:
                st.download_button(
                    "📥 Baixar DOC",
                    data=st.session_state.doc_bytes,
                    file_name="MeuRepertorio.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_doc"
                )
            
            # TXT
            if st.button("📄 Gerar TXT", key="gerar_txt"):
                with st.spinner("Gerando arquivo TXT..."):
                    st.session_state.txt_bytes = exportar_txt()
            
            if st.session_state.txt_bytes:
                st.download_button(
                    "📥 Baixar TXT",
                    data=st.session_state.txt_bytes,
                    file_name="MeuRepertorio.txt",
                    mime="text/plain",
                    key="download_txt"
                )
        
        with col2:
            st.subheader("📱 Kindle / Leitores")
            st.markdown("""
            **Como enviar para o Kindle:**
            1. Baixe o arquivo abaixo
            2. Envie por e-mail para seu `@kindle.com`
            3. Ou use o app Kindle
            """)
            
            if st.button("📱 Gerar TXT para Kindle", key="gerar_kindle"):
                with st.spinner("Gerando arquivo otimizado para Kindle..."):
                    st.session_state.kindle_bytes = exportar_txt_simples()
            
            if st.session_state.kindle_bytes:
                st.download_button(
                    "📥 Baixar para Kindle",
                    data=st.session_state.kindle_bytes,
                    file_name="Kindle_Repertorio.txt",
                    mime="text/plain",
                    key="download_kindle"
                )
        
        st.divider()
        
        st.subheader("📦 Exportar Completo (Todos Formatos)")
        if st.button("💾 Gerar ZIP com todos os formatos", key="gerar_zip"):
            with st.spinner("Gerando pacote completo..."):
                st.session_state.zip_bytes = exportar_completo_zip()
        
        if st.session_state.zip_bytes:
            st.download_button(
                "📥 Baixar ZIP Completo",
                data=st.session_state.zip_bytes,
                file_name="Repertorio_Completo.zip",
                mime="application/zip",
                key="download_zip"
            )
        
        st.divider()
        st.info("💡 **Dica para Kindle:** Baixe o arquivo TXT para Kindle e envie para seu e-mail do Kindle. O dispositivo converterá automaticamente.")
        st.success("✅ Todos os formatos disponíveis: PDF, DOC (Word), TXT e TXT otimizado para Kindle!")
